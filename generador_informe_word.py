import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px
import pandas as pd

# Directorio temporal para guardar las imágenes de los gráficos antes de pegarlas
TEMP_DIR = 'temp_images'
os.makedirs(TEMP_DIR, exist_ok=True)

def crear_informe_word_profesional(filepath, periodo, reuniones_raw, kpis, df_procesados, df_lic):
    doc = Document()
    
    # --- 1. PORTADA ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("Servicio Profesional Procesamiento de datos")
    run1.bold = True
    run1.font.size = Pt(18)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"INFORME DE ANÁLISIS PERÍODO {periodo.upper()} 2025")
    run2.font.size = Pt(14)
    doc.add_page_break()

    # --- 2. ACTA DE ENTREGA ---
    doc.add_heading('Acta de entrega', level=1)
    tabla_acta = doc.add_table(rows=0, cols=2)
    tabla_acta.style = 'Table Grid'
    
    datos_acta = [
        ["MATERIA", "SERVICIOS PROFESIONALES PROCESAMIENTO DE DATOS"],
        ["ID COMPRA", "2332-354-COT25"],
        ["LUGAR", "PUERTO MONTT. REGIÓN DE LOS LAGOS"],
        ["FECHA DE INICIO", "02/06/2025"],
        ["PROVEEDOR", "MARIA IGNACIA VARGAS TAPIA"],
        ["RUT", "20.794.682-6"]
    ]
    for key, val in datos_acta:
        row = tabla_acta.add_row().cells
        row[0].text = key
        row[1].text = val
    doc.add_paragraph("\n")

    # --- 3. METODOLOGÍA ---
    doc.add_heading('INICIO DEL SERVICIO Y METODOLOGÍA DE TRABAJO', level=2)
    doc.add_paragraph(
        "Con fecha 2 de junio de 2025, se dio inicio al servicio de procesamiento de datos destinado a la recolección "
        "y análisis de información del Departamento de Salud de Puerto Montt. [cite: 8, 9]"
    )

    # --- 4. REUNIONES ---
    doc.add_heading(f'Fechas Reunión {periodo}', level=2)
    t_reuniones = doc.add_table(rows=1, cols=2)
    t_reuniones.style = 'Table Grid'
    t_reuniones.rows[0].cells[0].text = "FECHA REUNIÓN"
    t_reuniones.rows[0].cells[1].text = "ESTADO"
    
    fechas = [f.strip() for f in reuniones_raw.split(',')]
    for f in fechas:
        row = t_reuniones.add_row().cells
        row[0].text = f
        row[1].text = "REALIZADA"

    doc.add_page_break()

    # --- 5. GRÁFICO 1: EFICIENCIA (TORTA) ---
    doc.add_heading('Rendimiento General de Requerimientos', level=2)
    
    # Crear gráfico de torta con Plotly
    img_pie_path = os.path.join(TEMP_DIR, "pie_eficiencia.png")
    fig_pie = px.pie(
        names=['Procesados', 'Cancelados'], 
        values=[kpis.get('req_procesados', 0), kpis.get('cancelados', 0)],
        color_discrete_sequence=['#2ecc71', '#e74c3c'],
        hole=0.4
    )
    fig_pie.write_image(img_pie_path, engine="kaleido")
    doc.add_picture(img_pie_path, width=Inches(4.5))
    
    doc.add_paragraph(
        f"Promedio de Efectividad: Se procesó un total de {kpis.get('total_neto', 0)} requerimientos limpios, "
        f"logrando una tasa del {kpis.get('eficiencia', 0)}% de conversión a Orden de Compra. [cite: 51]"
    )

    # --- 6. GRÁFICO 2: TOP UNIDADES (BARRAS) ---
    if df_procesados is not None and not df_procesados.empty:
        doc.add_page_break()
        doc.add_heading('Unidades Solicitantes de Mayor Impacto', level=2)
        
        # Procesar datos para el gráfico de barras
        top_unidades = df_procesados['unidad'].value_counts().head(5).sort_values(ascending=True)
        img_bar_path = os.path.join(TEMP_DIR, "bar_unidades.png")
        
        fig_bar = px.bar(
            x=top_unidades.values, 
            y=top_unidades.index, 
            orientation='h',
            labels={'x': 'Cantidad de OCs', 'y': 'Unidad'},
            color_discrete_sequence=['#3498db']
        )
        fig_bar.write_image(img_bar_path, engine="kaleido")
        doc.add_picture(img_bar_path, width=Inches(5.5))
        doc.add_paragraph("Análisis detallado de las unidades líderes en gasto estratégico. [cite: 75, 83]")

    # --- 7. LICITACIONES ---
    doc.add_page_break()
    doc.add_heading('Análisis de Licitaciones', level=2)
    t_lic = doc.add_table(rows=3, cols=2)
    t_lic.style = 'Table Grid'
    t_lic.rows[0].cells[0].text = "Total Licitaciones Gestionadas"
    t_lic.rows[0].cells[1].text = "9" # Dato fijo según informe de Noviembre [cite: 91]
    t_lic.rows[1].cells[0].text = "Licitaciones Adjudicadas"
    t_lic.rows[1].cells[1].text = f"{len(df_lic)}"
    t_lic.rows[2].cells[0].text = "Monto Total Adjudicado"
    t_lic.rows[2].cells[1].text = f"$ {df_lic['monto_adjudicado'].sum() if not df_lic.empty else 0:,.0f}"

    # --- 8. FIRMAS EN BLANCO ---
    doc.add_paragraph("\n\n\n\n\n") # Espacio para firma manual
    f_table = doc.add_table(rows=1, cols=2)
    f_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Celda 1: Tu espacio
    c1 = f_table.rows[0].cells[0]
    c1.text = "__________________________\nMARÍA IGNACIA VARGAS TAPIA\nProveedor\nPeríodo: " + periodo
    
    # Celda 2: Espacio Jefatura
    c2 = f_table.rows[0].cells[1]
    c2.text = "__________________________\nSRA. JOHANA\nJefa de Licitaciones\nDepto. Salud Puerto Montt"

    doc.save(filepath)
    return True